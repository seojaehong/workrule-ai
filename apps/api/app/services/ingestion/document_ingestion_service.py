from io import BytesIO
import logging
import os
import xml.etree.ElementTree as ET
import zlib
from zipfile import ZipFile

from docx import Document
from fastapi import HTTPException, UploadFile
import olefile
from pypdf import PdfReader

from app.domain.schemas.ingestion import ExtractedDocument
from app.services.ingestion.text_normalizer import normalize_extracted_text
from app.services.ingestion.vision_ocr_service import VisionOCRService


logger = logging.getLogger(__name__)


class DocumentIngestionService:
    _hwp_preview_stream = "PrvText"
    _hwp_paragraph_text_tag = 67
    _hwpx_preview_entry = "Preview/PrvText.txt"

    def __init__(self, *, ocr_service: VisionOCRService | None = None) -> None:
        self._ocr_service = ocr_service

    async def extract_text(self, file: UploadFile) -> ExtractedDocument:
        filename = file.filename or "uploaded-file"
        suffix = os.path.splitext(filename)[1].lower()
        content = await file.read()

        parser_name, extracted_text = await self._extract_with_parser(
            suffix=suffix,
            filename=filename,
            content=content,
        )
        normalized_text = normalize_extracted_text(extracted_text)
        if not normalized_text:
            raise HTTPException(status_code=422, detail="The uploaded file did not contain extractable text.")

        return ExtractedDocument(
            filename=filename,
            parser=parser_name,
            media_type=file.content_type,
            extracted_text=extracted_text,
            normalized_text=normalized_text,
            char_count=len(normalized_text),
            line_count=len(normalized_text.splitlines()),
        )

    async def _extract_with_parser(self, *, suffix: str, filename: str, content: bytes) -> tuple[str, str]:
        if suffix in {".txt", ".md"}:
            return "plain_text", content.decode("utf-8-sig", errors="ignore")
        if suffix == ".pdf":
            return await self._extract_pdf_text(content=content, filename=filename)
        if suffix == ".docx":
            return "docx", self._extract_docx_text(content)
        if suffix == ".hwpx":
            return "hwpx", self._extract_hwpx_text(content)
        if suffix == ".hwp":
            return "hwp", self._extract_hwp_text(content)
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type: {suffix or 'unknown'}. Use .txt, .md, .pdf, .docx, or .hwpx.",
        )

    async def _extract_pdf_text(self, *, content: bytes, filename: str) -> tuple[str, str]:
        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        extracted_text = "\n\n".join(page for page in pages if page.strip())
        if extracted_text.strip():
            return "pdf", extracted_text

        if self._ocr_service is not None:
            logger.info("Falling back to vision OCR for scanned PDF: %s", filename)
            ocr_text = await self._ocr_service.extract_pdf_text(content=content, filename=filename)
            return "pdf_vision_ocr", ocr_text

        raise HTTPException(
            status_code=422,
            detail=(
                "The uploaded PDF appears to be image-only scanned. Vision OCR is not configured yet. "
                "Set OPENAI_API_KEY to enable OCR, or upload an OCR-applied PDF or the original HWP/HWPX/DOCX file."
            ),
        )

    def _extract_docx_text(self, content: bytes) -> str:
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n".join(paragraphs)

    def _extract_hwpx_text(self, content: bytes) -> str:
        with ZipFile(BytesIO(content)) as archive:
            preview_text = self._extract_hwpx_preview_text(archive)
            if preview_text:
                return preview_text

            texts: list[str] = []
            section_names = sorted(
                name for name in archive.namelist() if name.startswith("Contents/section") and name.endswith(".xml")
            )
            if not section_names:
                raise HTTPException(status_code=422, detail="The .hwpx file did not contain readable section XML.")

            for section_name in section_names:
                root = ET.fromstring(archive.read(section_name))
                for node in root.iter():
                    if node.text and node.text.strip():
                        texts.append(node.text.strip())

        return "\n".join(texts)

    def _extract_hwpx_preview_text(self, archive: ZipFile) -> str | None:
        if self._hwpx_preview_entry not in archive.namelist():
            return None

        preview_bytes = archive.read(self._hwpx_preview_entry)
        for encoding in ("utf-8-sig", "utf-8"):
            try:
                preview_text = preview_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
            if preview_text.strip():
                return preview_text

        return None

    def _extract_hwp_text(self, content: bytes) -> str:
        with olefile.OleFileIO(BytesIO(content)) as archive:
            preview_text = self._extract_hwp_preview_text(archive)
            if preview_text:
                return preview_text

            section_names = self._list_hwp_section_names(archive)
            if not section_names:
                raise HTTPException(status_code=422, detail="The .hwp file did not contain readable body sections.")

            compressed = self._is_hwp_body_compressed(archive)
            texts = [
                self._extract_hwp_section_text(archive, section_name=section_name, compressed=compressed)
                for section_name in section_names
            ]
            return "\n".join(text for text in texts if text.strip())

    def _extract_hwp_preview_text(self, archive: olefile.OleFileIO) -> str | None:
        if not archive.exists(self._hwp_preview_stream):
            return None

        preview_bytes = archive.openstream(self._hwp_preview_stream).read()
        for encoding in ("utf-16", "utf-16le"):
            try:
                preview_text = preview_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
            if preview_text.strip():
                return preview_text

        return None

    def _list_hwp_section_names(self, archive: olefile.OleFileIO) -> list[str]:
        section_names = [
            "/".join(entry)
            for entry in archive.listdir(streams=True, storages=False)
            if len(entry) == 2 and entry[0] == "BodyText" and entry[1].startswith("Section")
        ]
        return sorted(section_names, key=self._hwp_section_sort_key)

    def _hwp_section_sort_key(self, section_name: str) -> int:
        return int(section_name.rsplit("Section", maxsplit=1)[1])

    def _is_hwp_body_compressed(self, archive: olefile.OleFileIO) -> bool:
        header = archive.openstream("FileHeader").read()
        if len(header) < 40:
            return False

        flags = int.from_bytes(header[36:40], "little")
        return bool(flags & 0x01)

    def _extract_hwp_section_text(
        self,
        archive: olefile.OleFileIO,
        *,
        section_name: str,
        compressed: bool,
    ) -> str:
        section_bytes = archive.openstream(section_name).read()
        if compressed:
            section_bytes = zlib.decompress(section_bytes, -15)

        texts: list[str] = []
        offset = 0
        while offset + 4 <= len(section_bytes):
            header = int.from_bytes(section_bytes[offset : offset + 4], "little")
            offset += 4

            tag_id = header & 0x3FF
            payload_size = (header >> 20) & 0xFFF
            if payload_size == 0xFFF:
                if offset + 4 > len(section_bytes):
                    break
                payload_size = int.from_bytes(section_bytes[offset : offset + 4], "little")
                offset += 4

            payload = section_bytes[offset : offset + payload_size]
            offset += payload_size

            if tag_id != self._hwp_paragraph_text_tag:
                continue

            decoded_text = payload.decode("utf-16le", errors="ignore").replace("\x00", "").strip()
            if decoded_text:
                texts.append(decoded_text)

        return "\n".join(texts)
